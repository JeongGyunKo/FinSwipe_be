from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, 16, True, (31, 78, 121))
    elif level == 2:
        set_font(run, 13, True, (31, 78, 121))
    elif level == 3:
        set_font(run, 11, True, (0, 0, 0))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 10)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 10, True, (255, 255, 255))
        shade_cell(cell, "1F4E79")

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, 9)
            if ri % 2 == 1:
                shade_cell(cell, "EBF3FB")

    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])

    doc.add_paragraph()
    return table


# ──────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# ── 표지 ──────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title_p.add_run("FinSwipe"), 36, True, (31, 78, 121))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(sub_p.add_run("AI 기반 미국 주식 뉴스 큐레이션 서비스"), 16, False, (89, 89, 89))

doc.add_paragraph()
doc.add_paragraph()

spec_p = doc.add_paragraph()
spec_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(spec_p.add_run("API 명세서 및 시스템 설계 문서"), 14, True, (0, 0, 0))

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(date_p.add_run("2026년 4월"), 11, False, (89, 89, 89))

doc.add_page_break()

# ── 1. 서비스 개요 ────────────────────────────
add_heading(doc, "1. 서비스 개요", 1)
add_body(
    doc,
    "FinSwipe는 미국 주식 시장 관련 뉴스를 AI가 자동으로 수집·분석하여 사용자에게 제공하는 서비스입니다. "
    "Finlight API를 통해 실시간으로 뉴스를 수집하고, Fine-tuning된 Grok 모델로 3줄 요약과 감성 분석을 수행합니다. "
    "사용자는 스와이프 UI로 뉴스를 빠르게 소비할 수 있으며, 번역은 Android 기기 내 ML Kit을 통해 온디바이스로 처리됩니다.",
)

doc.add_paragraph()

# ── 2. 시스템 아키텍처 ────────────────────────
add_heading(doc, "2. 시스템 아키텍처", 1)

add_table(
    doc,
    ["구분", "역할", "기술 스택", "배포"],
    [
        ["Backend", "뉴스 수집 / 저장 / REST API 제공", "FastAPI (Python)", "Zeabur"],
        ["GenAI", "3줄 요약, 감성 분석, XAI 추출", "Grok Fine-tuned", "Zeabur"],
        ["Frontend", "UI / 사용자 인터페이스", "Vercel (FE 프레임워크)", "Vercel"],
        ["Database", "기사 데이터 영구 저장", "Supabase (PostgreSQL)", "Supabase Cloud"],
        ["Android App", "뉴스 소비 / 온디바이스 번역", "Android + ML Kit Translation", "—"],
    ],
    [3.5, 5, 4.5, 3],
)

add_heading(doc, "데이터 흐름", 2)
for line in [
    "[Finlight API]",
    "      ↓  뉴스 수집 (15분 주기, APScheduler)",
    "[Backend — FastAPI / Zeabur]",
    "      ↓  기사 저장",
    "[Supabase DB — PostgreSQL]",
    "      ↓  분석 요청",
    "[GenAI Server — Grok Fine-tuned / Zeabur]",
    "      ↓  3줄 요약 + 감성 분석 + XAI (Attention Weight)",
    "[Supabase DB]",
    "      ↑  REST API 데이터 조회",
    "[Frontend — Vercel]",
    "      ↓  번역 (ML Kit, 온디바이스)",
    "[Android App]",
]:
    add_code(doc, line)

doc.add_paragraph()

# ── 3. 서버 정보 ──────────────────────────────
add_heading(doc, "3. 서버 정보", 1)
add_body(doc, "Backend Base URL: https://finswipe.zeabur.app")

doc.add_paragraph()

# ── 4. 공개 API 명세 ──────────────────────────
add_heading(doc, "4. 공개 API 명세", 1)
add_body(
    doc,
    "아래 엔드포인트는 별도 인증 없이 호출 가능한 공개 API입니다. Rate Limit: 30회/분 (GenAI 관련 10회/분)",
)

# 4-1
add_heading(doc, "4.1  서버 상태 확인", 2)
add_code(doc, "GET /health")
add_body(doc, "서버, 데이터베이스, GenAI 서버의 전체 상태를 반환합니다.")
add_table(
    doc,
    ["필드", "타입", "설명"],
    [
        ["status", "string", "ok | degraded — 전체 서버 상태"],
        ["db", "string", "ok | error — Supabase 연결 상태"],
        ["genai", "string", "ok | suspended | offline | error — GenAI 서버 상태"],
    ],
    [3, 3, 10],
)
add_heading(doc, "응답 예시", 3)
add_code(doc, '{ "status": "ok", "db": "ok", "genai": "ok" }')

# 4-2
add_heading(doc, "4.2  최신 뉴스 조회", 2)
add_code(doc, "GET /news/latest")
add_body(doc, "감성 분석, 3줄 요약, XAI가 포함된 최신 뉴스 기사를 반환합니다.")
add_table(
    doc,
    ["파라미터", "타입", "기본값", "설명"],
    [
        ["limit", "int", "20", "조회 개수 (최대 100)"],
        ["offset", "int", "0", "페이지네이션 오프셋"],
    ],
    [3, 3, 3, 7],
)
add_heading(doc, "응답 예시", 3)
for line in [
    "{",
    '  "count": 20, "offset": 0,',
    '  "data": [{',
    '    "id": "uuid",',
    '    "headline": "Apple beats earnings expectations",',
    '    "summary": "...",',
    '    "summary_3lines": ["line1", "line2", "line3"],',
    '    "source_url": "https://...",',
    '    "image_url": "https://...",',
    '    "published_at": "2026-04-13T00:00:00Z",',
    '    "tickers": ["AAPL"],',
    '    "ticker_names": [{ "ticker": "AAPL", "corp": "Apple Inc.", "ko": "애플" }],',
    '    "sentiment_label": "positive",',
    '    "sentiment_score": 85.0,',
    '    "xai": { "highlights": [ ... ] },',
    '    "is_mixed": false,',
    '    "categories": ["technology"], "countries": ["US"]',
    "  }]",
    "}",
]:
    add_code(doc, line)

# 4-3
add_heading(doc, "4.3  뉴스 검색", 2)
add_code(doc, "GET /news/search?q={검색어}")
add_body(doc, "한국어 회사명, 영문 회사명, 티커 심볼 모두 검색 가능합니다. (예: 애플, Apple, AAPL)")
add_table(
    doc,
    ["파라미터", "타입", "필수", "설명"],
    [
        ["q", "string", "필수", "검색어 (예: 애플, Apple, AAPL)"],
        ["limit", "int", "선택", "기본 20, 최대 100"],
        ["offset", "int", "선택", "기본 0"],
    ],
    [3, 3, 2, 8],
)
add_heading(doc, "응답 예시", 3)
for line in [
    "{",
    '  "count": 5, "offset": 0,',
    '  "query": "애플",',
    '  "matched_tickers": ["AAPL"],',
    '  "data": [ /* 기사 객체 배열 (4.2와 동일) */ ]',
    "}",
]:
    add_code(doc, line)

# 4-4
add_heading(doc, "4.4  티커 목록 조회", 2)
add_code(doc, "GET /news/tickers")
add_body(
    doc,
    "지원하는 미국 주식 티커 전체 목록(134개)을 반환합니다. 프론트엔드 검색 자동완성에 활용합니다.",
)
add_heading(doc, "응답 예시", 3)
for line in [
    "{",
    '  "count": 134,',
    '  "data": [',
    '    { "ticker": "AAPL", "corp": "Apple Inc.", "ko": "애플" },',
    '    { "ticker": "MSFT", "corp": "Microsoft Corp.", "ko": "마이크로소프트" }',
    "  ]",
    "}",
]:
    add_code(doc, line)

# 4-5
add_heading(doc, "4.5  GenAI 서버 상태 확인", 2)
add_code(doc, "GET /news/genai/health")
add_table(
    doc,
    ["필드", "값"],
    [["status", "ok | suspended | offline | error"]],
    [4, 12],
)

doc.add_paragraph()

# ── 5. 기사 데이터 필드 정의 ───────────────────
add_heading(doc, "5. 기사 데이터 필드 정의", 1)
add_table(
    doc,
    ["필드", "타입", "설명"],
    [
        ["id", "string (UUID)", "기사 고유 ID"],
        ["headline", "string", "기사 제목 (영문 원문)"],
        ["summary", "string", "기사 짧은 요약 (영문)"],
        ["summary_3lines", "string[]", "GenAI가 생성한 3줄 요약 (영문)"],
        ["source_url", "string", "원문 기사 URL"],
        ["image_url", "string | null", "썸네일 이미지 URL"],
        ["published_at", "string (ISO 8601)", "기사 발행 시각"],
        ["tickers", "string[]", '관련 종목 티커 (예: ["AAPL", "MSFT"])'],
        ["ticker_names", "object[]", "티커별 한국어/영문 회사명"],
        ["sentiment_label", "string | null", "감성 레이블 (positive | negative | neutral)"],
        ["sentiment_score", "float | null", "감성 점수 (-100 ~ 100)"],
        ["xai", "object | null", "감성 판단 근거 (Attention Weight 기반)"],
        ["is_mixed", "bool | null", "복합 감성 여부 (긍·부정 혼재)"],
        ["categories", "string[]", "기사 카테고리"],
        ["countries", "string[]", "관련 국가 코드"],
    ],
    [4, 4, 8],
)

# ── 6. XAI 필드 상세 ──────────────────────────
add_heading(doc, "6. XAI 필드 상세", 1)
add_body(
    doc,
    "감성 분석 결과의 판단 근거를 Attention Weight 기반으로 제공합니다. "
    "어떤 문구가 긍정/부정 판단에 기여했는지 프론트엔드 하이라이트 표시에 활용합니다.",
)
for line in [
    '"xai": {',
    '  "highlights": [',
    "    {",
    '      "text_snippet": "beat earnings expectations",',
    '      "contribution_direction": "positive",',
    '      "importance_score": 0.92',
    "    },",
    "    {",
    '      "text_snippet": "supply chain concerns",',
    '      "contribution_direction": "negative",',
    '      "importance_score": 0.45',
    "    }",
    "  ]",
    "}",
]:
    add_code(doc, line)

add_table(
    doc,
    ["필드", "타입", "설명"],
    [
        ["text_snippet", "string", "감성에 영향을 준 핵심 문구"],
        ["contribution_direction", "string", "기여 방향 (positive | negative)"],
        ["importance_score", "float", "중요도 (0.0 ~ 1.0)"],
    ],
    [5, 5, 6],
)

# ── 7. GenAI 처리 흐름 ────────────────────────
add_heading(doc, "7. GenAI 처리 흐름", 1)
add_body(
    doc,
    "Backend에서 기사를 수집·저장한 뒤 GenAI 서버에 비동기로 제출합니다. "
    "GenAI 서버는 Fine-tuning된 Grok 모델을 사용하여 아래 3단계 분석을 수행합니다.",
)
add_table(
    doc,
    ["단계", "항목", "설명"],
    [
        ["1", "3줄 요약 생성", "기사 핵심 내용을 영문 3문장으로 요약"],
        ["2", "감성 분석", "positive / negative / neutral 분류 + 점수(-100~100)"],
        ["3", "XAI 근거 추출", "Attention Weight 기반으로 판단 근거 문구 추출"],
    ],
    [2, 5, 9],
)
for line in [
    "기사 수집 완료",
    "      ↓",
    "GenAI 서버 제출 (Grok Fine-tuned)",
    "      ↓",
    "┌──────────────────────────────────┐",
    "│  1. 3줄 요약 생성                 │",
    "│  2. 감성 분석 (pos / neg / neu)   │",
    "│  3. XAI 근거 추출                 │",
    "│     (Attention Weight 기반)       │",
    "└──────────────────────────────────┘",
    "      ↓",
    "Supabase DB 저장 → API 응답에 포함",
]:
    add_code(doc, line)

doc.add_paragraph()

# ── 8. 번역 처리 방식 ────────────────────────
add_heading(doc, "8. 번역 처리 방식", 1)
add_body(
    doc,
    "번역은 서버가 아닌 클라이언트(Android 앱)에서 온디바이스로 처리합니다. "
    "서버 비용 절감 및 응답 속도 향상을 위한 설계이며, API 키나 토큰이 전혀 필요하지 않습니다.",
)
add_table(
    doc,
    ["항목", "내용"],
    [
        ["라이브러리", "Google ML Kit Translation"],
        ["처리 방식", "온디바이스 (서버 통신 없음)"],
        ["비용", "무료 (토큰/API 키 불필요)"],
        ["번역 대상", "headline, summary_3lines 영→한"],
        ["특징", "최초 실행 시 언어팩 다운로드 (~30MB), 이후 오프라인 동작 가능"],
    ],
    [4, 12],
)
add_body(doc, "백엔드 API는 모든 텍스트를 영문 원문으로 제공하며, 번역은 앱에서 처리합니다.")

doc.add_paragraph()

# ── 9. Rate Limit ─────────────────────────────
add_heading(doc, "9. Rate Limit", 1)
add_table(
    doc,
    ["구분", "제한"],
    [
        ["공개 엔드포인트 (/news/latest, /search, /tickers)", "30회 / 분"],
        ["GenAI 관련 (/news/genai/health)", "10회 / 분"],
    ],
    [10, 6],
)
add_body(doc, "제한 초과 시 429 Too Many Requests 응답을 반환합니다.")

doc.save("c:/dev/FinSwipe_be/FinSwipe_API_명세서.docx")
print("저장 완료")
